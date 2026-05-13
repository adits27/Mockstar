import io
import re

from fastapi import HTTPException, UploadFile


async def extract_text(file: UploadFile) -> str:
    data = await file.read()
    ct = file.content_type or ""
    name = file.filename or ""

    if ct == "application/pdf" or name.lower().endswith(".pdf"):
        return _parse_pdf(data)
    if ct in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ) or name.lower().endswith((".docx", ".doc")):
        return _parse_docx(data)

    raise HTTPException(status_code=415, detail="Unsupported file type. Upload a PDF or DOCX.")


def _parse_pdf(data: bytes) -> str:
    import pypdf  # deferred so the rest of the app works without it in tests

    reader = pypdf.PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text)
    return _normalise(" ".join(parts))


def _parse_docx(data: bytes) -> str:
    import docx  # python-docx

    doc = docx.Document(io.BytesIO(data))
    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Also pull text from tables (common in formatted CVs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    return _normalise("\n".join(parts))


def _normalise(text: str) -> str:
    text = re.sub(r"\r\n|\r", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
